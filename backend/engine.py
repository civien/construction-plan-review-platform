#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
engine.py — 施工方案审查核心引擎（可被 FastAPI 直接调用）
  - extract_docx : docx -> 章节骨架 + 段落 + 表格 -> full_text.txt / structure.json
  - analyze      : 结构完整性 / 规范版本 / 跨章节一致性（机械层，读 JSON 规则）
  - build_annotated_docx : 把 findings 以 Word 批注形式写入方案 docx 副本
  - build_markdown_report : 生成 Markdown 审查报告
规则与执行解耦：analyze / build 都接收 rules(JSON) 与 directory_template(文本)。
"""
import os, re, json, datetime, zipfile
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
RT_COMMENTS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
PKG_REL = 'http://schemas.openxmlformats.org/package/2006/relationships'

CHAP_RE = re.compile(r'^第[一二三四五六七八九十百\d]+章[\s、．.]*')
SECTION_RE = re.compile(r'^(第[一二三四五六七八九十]+节|[\u4e00-\u9fa5]+[、．.])')
NUM_RE = re.compile(r'^([一二三四五六七八九十\d]+[、．.\s][\u4e00-\u9fa5])')


def _detect_level(text, style):
    t = text.strip()
    if CHAP_RE.match(t):
        return 1
    st = (style or '')
    m = re.match(r'Heading\s*(\d+)', st)
    if m:
        return int(m.group(1))
    if 'Heading' in st or '标题' in st:
        return 2
    if SECTION_RE.match(t) or NUM_RE.match(t):
        return 3
    return 0


def _is_chapter(text, style):
    t = text.strip()
    if CHAP_RE.match(t):
        return True
    if style and ('Heading 1' in style or '标题 1' in style):
        return True
    return False


# ---------------- 1. 抽取 ----------------
def extract_docx(input_path, out_dir):
    """返回 (full_text:str, structure:dict)，并写出 full_text.txt / structure.json"""
    os.makedirs(out_dir, exist_ok=True)
    doc = Document(input_path)

    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        style = p.style.name if p.style else ''
        paragraphs.append({
            'idx': i, 'style': style, 'text': text,
            'level': _detect_level(text, style),
            'is_chapter': _is_chapter(text, style),
        })

    tables = []
    for ti, tb in enumerate(doc.tables):
        rows = [[c.text.strip() for c in row.cells] for row in tb.rows]
        tables.append({'idx': ti, 'rows': rows, 'n': len(rows),
                       'cols': len(rows[0]) if rows else 0})

    # 表格在 body 中的顺序 -> 其前最近段落 idx
    body = doc.element.body
    pidx = -1
    tbl_idx = -1
    para_to_tbl_before = {}
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            pidx += 1
        elif child.tag == qn('w:tbl'):
            tbl_idx += 1
            para_to_tbl_before[tbl_idx] = pidx

    chapters = []
    cur = None

    def chapter_of_para(pidx):
        best = None
        for ch in chapters:
            if ch['para_idx'] <= pidx:
                best = ch
        return best

    for p in paragraphs:
        if p['is_chapter']:
            cur = {'title': p['text'].strip(), 'para_idx': p['idx'],
                   'items': [], 'tables': []}
            chapters.append(cur)
        else:
            if cur is None:
                if not chapters or chapters[0].get('title') != '__前置__':
                    chapters.insert(0, {'title': '__前置__', 'para_idx': -1,
                                        'items': [], 'tables': []})
                    cur = chapters[0]
            cur['items'].append(p)

    for ti, tb in enumerate(tables):
        pbefore = para_to_tbl_before.get(ti, 0)
        ch = chapter_of_para(pbefore) if pbefore is not None else (chapters[-1] if chapters else None)
        if ch is None:
            ch = {'title': '__前置__', 'para_idx': -1, 'items': [], 'tables': []}
            chapters.insert(0, ch)
        ch['tables'].append(tb)

    ft = []
    ft.append(f"# 文档结构提取: {os.path.basename(input_path)}")
    ft.append(f"段落总数: {len(paragraphs)}  表格总数: {len(tables)}")
    ft.append("=" * 70)
    for ch in chapters:
        ft.append("")
        ft.append(f"\n## {ch['title']}")
        for p in ch['items']:
            t = p['text'].strip()
            if not t:
                continue
            indent = '  ' * max(0, (p['level'] - 1))
            ft.append(f"{indent}- {t}")
        for tb in ch['tables']:
            ft.append(f"\n  [表格 {tb['idx'] + 1}] ({tb['n']}行 x {tb['cols']}列)")
            for r in tb['rows']:
                ft.append("    | " + " | ".join(r))
            ft.append("")
    full_text = "\n".join(ft)

    structure = {
        'meta': {'file': os.path.basename(input_path),
                 'paragraphs': len(paragraphs), 'tables': len(tables)},
        'chapters': [{
            'title': c['title'], 'para_idx': c['para_idx'],
            'n_items': len(c['items']), 'n_tables': len(c['tables']),
            'table_idx': [t['idx'] for t in c['tables']],
        } for c in chapters],
        'paragraphs': paragraphs,
        'tables': tables,
    }

    with open(os.path.join(out_dir, 'full_text.txt'), 'w', encoding='utf-8') as f:
        f.write(full_text)
    with open(os.path.join(out_dir, 'structure.json'), 'w', encoding='utf-8') as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)
    return full_text, structure


# ---------------- 2. 机械分析 ----------------
def _base_code(s):
    s = s.replace(' ', '').replace('-', '').replace('/', '')
    m = re.match(r'([A-Za-z]+)(\d+)', s)
    return m.group(1) + m.group(2) if m else s


def analyze(full_text, structure, rules, template_text=''):
    lines = full_text.split('\n')

    # 2.1 结构完整性
    exp_chapters = re.findall(r'^#\s*第[一二三四五六七八九十]+章\s*.+$', template_text or '', re.M)
    doc_chapters = [c['title'] for c in structure['chapters'] if c['title'] != '__前置__']

    def norm_title(t):
        return re.sub(r'\s+', '', t)

    exp_ch_norm = [norm_title(re.sub(r'^#\s*', '', x)) for x in exp_chapters]
    doc_ch_norm = [norm_title(x) for x in doc_chapters]
    missing_chapters = [x for x in exp_ch_norm if x not in doc_ch_norm]

    # 2.2 规范版本核查
    norm_versions = rules.get('norm_versions', {}) or {}
    valid_by_base = {}
    for code, info in norm_versions.items():
        valid_by_base[_base_code(code)] = {'code': code,
                                           'name': info.get('name', ''),
                                           'note': info.get('note', '')}
    code_pat = re.compile(
        r'(GB|JGJ|JGJ/T|CECS|DBJ|DGJ|11SG|03SG|04SG|05SG|06SG|07SG|08SG|09SG|10SG|12SG|13SG|14SG|15SG|16SG|17SG|18SG|19SG|20SG|21SG)\s*[\u4e00-\u9fa5]?\s*[-\s]?\s*(\d{3,4})\s*[-\s]?\s*(\d{2,4})?')
    book_pat = re.compile(
        r'《([^》]+)》[^《\n]{0,30}?(GB|JGJ|JGJ/T|CECS|DBJ|DGJ|11SG|03SG|04SG|05SG|06SG|07SG|08SG|09SG|10SG|12SG|13SG|14SG|15SG|16SG|17SG|18SG|19SG|20SG|21SG)\s*[-\s]?\s*(\d{3,4})\s*[-\s]?\s*(\d{2,4})?')
    found_codes = {}
    for m in book_pat.finditer(full_text):
        name, prefix, num, yr = m.groups()
        key = f"{prefix} {num}"
        found_codes.setdefault(key, {'years': set(), 'names': set()})
        if yr:
            found_codes[key]['years'].add(yr)
        found_codes[key]['names'].add(name)
    for m in code_pat.finditer(full_text):
        prefix, num, yr = m.groups()
        key = f"{prefix} {num}"
        found_codes.setdefault(key, {'years': set(), 'names': set()})
        if yr:
            found_codes[key]['years'].add(yr)

    norm_findings = []
    for key, info in found_codes.items():
        kb = _base_code(key)
        if kb in valid_by_base:
            v = valid_by_base[kb]
            valid_yr = re.search(r'(\d{4})', v['code'])
            valid_yr = valid_yr.group(1) if valid_yr else ''
            yrs = info['years']
            if yrs:
                for y in sorted(yrs):
                    if valid_yr and y != valid_yr:
                        anchor = f"{key}-{y}"
                        norm_findings.append({
                            'code': key, 'anchor': anchor, 'severity': '中',
                            'text': f"【规范版本】文档出现 {anchor}（{','.join(list(info['names'])[:1])}），"
                                    f"与有效版本 {v['code']}（{v['name']}）年份不一致，请核对。"})
                    else:
                        norm_findings.append({
                            'code': key, 'anchor': f"{key} {y}", 'severity': '轻',
                            'text': f"【规范版本】{key} {y} 与有效版本 {v['code']} 一致（{v['name']}）。"})
            else:
                norm_findings.append({
                    'code': key, 'anchor': key, 'severity': '待核',
                    'text': f"【规范版本】{key}（{','.join(list(info['names'])[:1]) or '?'}）未标注年份，"
                            f"有效版本应为 {v['code']}，请确认。"})
        else:
            norm_findings.append({
                'code': key, 'anchor': key, 'severity': '待核',
                'text': f"【规范版本】{key} {sorted(info['years']) or '无年份'}（{','.join(list(info['names'])[:1]) or '?'}）"
                        f"未在规则规范表，需人工确认有效性。"})

    # 2.3 跨章节取值提取
    consistency = {}
    for item in rules.get('cross_chapter_checks', []) or []:
        key = item['key']
        kws = item.get('keywords', [])
        occ = []
        for i, ln in enumerate(lines):
            if any(k in ln for k in kws):
                occ.append({'line': i, 'text': ln.strip()})
        consistency[key] = occ

    # 2.4 严重缺陷 & 清单
    severe = rules.get('severe_defects', []) or []
    checklist_count = len(rules.get('checklist', []) or [])

    # 2.5 快速矛盾检测（2.0 新增）：关键词命中提取，供 LLM 做语义矛盾判定
    quick_hits = []
    for item in rules.get('quick_conflict_checks', []) or []:
        kws = item.get('keywords', []) or []
        occ = []
        for i, ln in enumerate(lines):
            if any(k in ln for k in kws):
                occ.append({'line': i, 'text': ln.strip()})
        quick_hits.append({
            'item': item.get('item', ''),
            'check': item.get('check', ''),
            'keywords': kws,
            'hit_count': len(occ),
            'occ': occ[:20],
        })

    skeleton = {
        'missing_chapters': missing_chapters,
        'doc_chapters': doc_chapters,
        'norm_findings': norm_findings,
        'consistency': consistency,
        'quick_hits': quick_hits,
        'severe': severe,
        'checklist_count': checklist_count,
    }
    return skeleton


# ---------------- 3. 生成批注 docx ----------------
def _qn(tag):
    return '{%s}%s' % (W, tag)


def _find_para(doc_root, substr):
    for p in doc_root.iter(_qn('p')):
        tx = ''.join(t.text or '' for t in p.iter(_qn('t')))
        if substr in tx:
            return p
    return None


def _insert_comment(paragraph, cid):
    runs = paragraph.findall(_qn('r'))
    if not runs:
        r = etree.SubElement(paragraph, _qn('r'))
        runs = [r]
    first, last = runs[0], runs[-1]
    crs = etree.Element(_qn('commentRangeStart'))
    crs.set(_qn('id'), str(cid))
    first.addprevious(crs)
    cre = etree.Element(_qn('commentRangeEnd'))
    cre.set(_qn('id'), str(cid))
    last.addnext(cre)
    ref = etree.Element(_qn('r'))
    rpr = etree.SubElement(ref, _qn('rPr'))
    rst = etree.SubElement(rpr, _qn('rStyle'))
    rst.set(_qn('val'), 'CommentReference')
    cref = etree.SubElement(ref, _qn('commentReference'))
    cref.set(_qn('id'), str(cid))
    cre.addnext(ref)


def _add_run(p, txt):
    r = etree.SubElement(p, _qn('r'))
    rpr = etree.SubElement(r, _qn('rPr'))
    rst = etree.SubElement(rpr, _qn('rStyle'))
    rst.set(_qn('val'), 'CommentText')
    t = etree.SubElement(r, _qn('t'))
    t.text = txt
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _make_comment_xml(comments, author='方案审查'):
    root = etree.Element(_qn('comments'))
    date = datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')
    for cid, sev, text in comments:
        c = etree.SubElement(root, _qn('comment'))
        c.set(_qn('id'), str(cid))
        c.set(_qn('author'), author)
        c.set(_qn('date'), date)
        c.set(_qn('initials'), '审')
        p = etree.SubElement(c, _qn('p'))
        _add_run(p, f"[{sev}] ")
        for i, line in enumerate(text.split('\n')):
            if i > 0:
                br = etree.SubElement(p, _qn('r'))
                etree.SubElement(br, _qn('br'))
            _add_run(p, line)
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def build_annotated_docx(input_docx, output_docx, findings, author='方案审查'):
    """findings: list of {anchor, severity, text}"""
    zin = zipfile.ZipFile(input_docx, 'r')
    root = etree.fromstring(zin.read('word/document.xml'))
    body = root.find(_qn('body'))

    comments = []
    cid = 0
    not_found = []
    for item in findings:
        substr = item.get('anchor', '')
        sev = item.get('severity', '中')
        text = item.get('text', '')
        p = _find_para(root, substr)
        if p is None:
            not_found.append(item)
            continue
        cid += 1
        _insert_comment(p, cid)
        comments.append((cid, sev, text))

    # 未锚定：追加到文末
    if not_found:
        for item in not_found:
            cid += 1
            p = etree.SubElement(body, _qn('p'))
            r = etree.SubElement(p, _qn('r'))
            t = etree.SubElement(r, _qn('t'))
            t.text = f"（审查批注-未定位原段落）{item.get('anchor', '')}"
            _insert_comment(p, cid)
            comments.append((cid, item.get('severity', '中'), item.get('text', '')))

    comments_bytes = _make_comment_xml(comments, author)

    ct = etree.fromstring(zin.read('[Content_Types].xml'))
    if not any(ov.get('PartName') == '/word/comments.xml'
               for ov in ct.findall('{%s}Override' % CT)):
        ov = etree.SubElement(ct, '{%s}Override' % CT)
        ov.set('PartName', '/word/comments.xml')
        ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')

    rels = etree.fromstring(zin.read('word/_rels/document.xml.rels'))
    maxid = 0
    for rel in rels.findall('{%s}Relationship' % PKG_REL):
        m = re.match(r'rId(\d+)', rel.get('Id') or '')
        if m:
            maxid = max(maxid, int(m.group(1)))
    rel = etree.SubElement(rels, '{%s}Relationship' % PKG_REL)
    rel.set('Id', f'rId{maxid + 1}')
    rel.set('Type', RT_COMMENTS)
    rel.set('Target', 'comments.xml')

    styles = etree.fromstring(zin.read('word/styles.xml'))
    style_ids = {s.get('{%s}styleId' % W) for s in styles.findall('{%s}style' % W)}
    for need in ('CommentReference', 'CommentText'):
        if need not in style_ids:
            s = etree.SubElement(styles, '{%s}style' % W)
            s.set('{%s}type' % W, 'character')
            s.set('{%s}styleId' % W, need)
            s.set('{%s}default' % W, 'false')
            nm = etree.SubElement(s, '{%s}name' % W)
            nm.set('{%s}val' % W, need)

    with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == 'word/document.xml':
                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
            elif item == '[Content_Types].xml':
                data = etree.tostring(ct, xml_declaration=True, encoding='UTF-8', standalone=True)
            elif item == 'word/_rels/document.xml.rels':
                data = etree.tostring(rels, xml_declaration=True, encoding='UTF-8', standalone=True)
            elif item == 'word/styles.xml':
                data = etree.tostring(styles, xml_declaration=True, encoding='UTF-8', standalone=True)
            zout.writestr(item, data)
        zout.writestr('word/comments.xml', comments_bytes)
    zin.close()
    return {'comments': len(comments), 'not_found': len(not_found)}


# ---------------- 4. Markdown 报告 ----------------
def build_markdown_report(meta, skeleton, findings, used_llm):
    RATING_MAP = {
        'excellent': ('优秀方案', '✅', '方案质量优秀，问题极少，可直接通过审查。'),
        'pass': ('合格方案', '🔵', '方案基本符合要求，存在少量可优化项，建议修改后通过。'),
        'warn': ('待优化方案', '🟠', '方案存在较明显缺陷或矛盾，需重点修改后重新审查。'),
        'fail': ('不合格方案', '❌', '方案存在严重缺陷或大量问题，必须全面修订后重新提交。'),
    }
    rating = meta.get('rating', '')
    rlabel, ricon, rdesc = RATING_MAP.get(rating, ('未评级', '⚪', ''))

    lines = []
    lines.append(f"# 施工方案审查报告（{meta.get('type_name', '')}）")
    lines.append("")
    if rating:
        lines.append(f"> **综合评级：{ricon} {rlabel}** — {rdesc}")
        lines.append("")
    lines.append(f"- 文件：{meta.get('filename', '')}")
    lines.append(f"- 审查时间：{meta.get('created_at', '')}")
    lines.append(f"- 引擎：{'大模型判读' if used_llm else '机械分析（未接入大模型）'}")
    lines.append(f"- 批注条数：{len(findings)}")
    lines.append("")
    lines.append("## 一、问题清单（按严重度）")
    lines.append("")
    order = {'严重': 0, '中': 1, '轻': 2, '待核': 3, '总评': 4}
    for f in sorted(findings, key=lambda x: order.get(x.get('severity', '轻'), 9)):
        lines.append(f"- **[{f.get('severity', '轻')}]** {f.get('text', '').replace(chr(10), ' / ')}")
    lines.append("")
    lines.append("## 二、结构完整性")
    if skeleton.get('missing_chapters'):
        lines.append("缺失章节：" + "；".join(skeleton['missing_chapters']))
    else:
        lines.append("章节齐全，无缺失。")
    lines.append("")
    lines.append("## 三、规范版本核查")
    for nf in skeleton.get('norm_findings', []):
        lines.append(f"- {nf.get('text', '')}")
    lines.append("")
    lines.append("## 四、快速矛盾检测（2.0 新增）")
    for q in skeleton.get('quick_hits', []):
        lines.append(f"- **{q.get('item', '')}**｜命中 {q.get('hit_count', 0)} 处｜检查要点：{q.get('check', '')}")
        for o in q.get('occ', [])[:3]:
            lines.append(f"    - 第{o['line']}行：{o['text']}")
    lines.append("")
    lines.append("## 五、严重缺陷清单（一票否决项）")
    for s in skeleton.get('severe', []):
        lines.append(f"- {s.get('no', '')}. {s.get('desc', '')}")
    lines.append("")
    lines.append("> 说明：本报告由施工方案审查平台自动生成，最终判定以专家复核为准。")
    return "\n".join(lines)


if __name__ == '__main__':
    import sys
    print(analyze.__doc__)
